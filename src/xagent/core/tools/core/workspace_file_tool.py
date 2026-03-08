"""
Core workspace file operations for xagent

This module provides the core file operation logic that works with workspace instances.
It focuses on pure file operations without tool framework dependencies.
"""

import asyncio
import csv
import logging
import os
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from pydantic import BaseModel

from ...workspace import TaskWorkspace
from .document_parser import DocumentCapabilities, DocumentParseArgs, parse_document
from .file_tool import (
    DEFAULT_MAX_FILE_SIZE,
    PREVIEW_LINES,
    EditOperation,
    EditResult,
    _get_file_preview_from_file,
    _is_binary_by_mime,
    _is_binary_file,
)

logger = logging.getLogger(__name__)


def is_document_file(file_path: str) -> bool:
    """Check if file is a document format that requires special parsing."""
    document_extensions = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".md"}
    return Path(file_path).suffix.lower() in document_extensions


def extract_text_from_document(file_path: str) -> str:
    """Extract text content from a document file using document parser with fallback."""
    file_ext = Path(file_path).suffix.lower()

    # Try different parsers in order of preference
    parsers_to_try = []

    if file_ext == ".pdf":
        parsers_to_try = ["deepdoc", "unstructured", "pypdf", "pdfplumber", "pymupdf"]
    elif file_ext == ".docx":
        parsers_to_try = ["deepdoc"]  # Only DeepDoc supports DOCX
    elif file_ext in [".xlsx", ".xls", ".csv"]:
        parsers_to_try = ["deepdoc"]
    elif file_ext == ".md":
        parsers_to_try = ["deepdoc"]
    else:
        parsers_to_try = ["deepdoc"]

    last_error: Exception | None = None
    for parser_name in parsers_to_try:
        try:
            logger.debug(f"Trying to parse {file_path} with {parser_name}")

            # Run the async document parser in sync context
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

            try:
                parse_args = DocumentParseArgs(
                    file_path=file_path,
                    parser_name=parser_name,
                    capabilities=DocumentCapabilities(
                        capability_text=True,
                        capability_figure=True,
                        requires_full_text_result=True,
                        requires_segmented_result=False,
                        use_local_parser=True,
                    ),
                )

                result = loop.run_until_complete(parse_document(parse_args))

                # Extract text from all segments
                text_parts = []
                if result.text_segments:
                    text_parts.extend(
                        [segment.text for segment in result.text_segments]
                    )

                # Add table content if available
                if result.tables:
                    for table in result.tables:
                        if table.html:
                            text_parts.append(f"Table:\n{table.html}")

                # Add figure captions if available
                if result.figures:
                    for figure in result.figures:
                        if figure.text:
                            text_parts.append(f"Figure: {figure.text}")

                text_content = "\n\n".join(text_parts) if text_parts else ""

                if text_content.strip():
                    logger.info(f"Successfully parsed {file_path} with {parser_name}")
                    return text_content
                else:
                    logger.warning(
                        f"Parser {parser_name} returned empty content for {file_path}"
                    )
                    last_error = ValueError(
                        f"Parser {parser_name} returned empty content"
                    )
                    continue

            finally:
                loop.close()

        except Exception as e:
            logger.warning(f"Failed to parse {file_path} with {parser_name}: {e}")
            last_error = e
            continue

    # All parsers failed, try basic python-docx fallback for DOCX
    if file_ext == ".docx":
        try:
            logger.info(f"Attempting python-docx fallback for {file_path}")
            import docx

            doc: Any = docx.Document(file_path)

            # Extract all paragraphs and tables
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    text_parts.append("Table:\n" + "\n".join(table_text))

            content = "\n\n".join([text for text in text_parts if text.strip()])
            if content.strip():
                logger.info(
                    f"Successfully extracted text from {file_path} using python-docx fallback"
                )
                return content

        except Exception as docx_error:
            logger.warning(
                f"python-docx fallback also failed for {file_path}: {docx_error}"
            )

    # All parsers and fallbacks failed, raise the last error
    logger.error(f"All parsers failed for {file_path}. Last error: {last_error}")
    raise ValueError(
        f"Unable to parse document {file_path} with any available parser. Last error: {last_error}"
    )


class FileInfo(BaseModel):
    """File information model"""

    name: str
    path: str
    size: int
    is_file: bool
    is_dir: bool
    modified_time: float
    encoding: Optional[str] = None


class ListFilesResult(BaseModel):
    """List files result model"""

    files: List[FileInfo]
    total_count: int
    current_path: str


class WorkspaceFileOperations:
    """
    Core workspace file operations.

    This class provides pure file operation logic without tool framework dependencies.
    It works with workspace instances to ensure operations are restricted to workspace boundaries.
    """

    def __init__(self, workspace: TaskWorkspace):
        self.workspace = workspace

    def read_file(self, file_path: str, encoding: str = "utf-8") -> str:
        """
        Read the content of a text file from the workspace.

        This function can only read text files. Binary files (like images, executables) will be rejected.
        Document files (PDF, DOCX, XLSX, etc.) will be automatically parsed to extract text content.
        For files larger than 1MB, a preview with the first 100 lines will be returned instead.

        The file path is searched in this order:
        1. Input directory (for user-uploaded files)
        2. Output directory (for generated files)

        Args:
            file_path: Path to the file (relative to workspace directories)
            encoding: Text encoding format (default: utf-8). Most text files use utf-8.

        Returns:
            - For small files (≤1MB): The complete file content as text
            - For large files (>1MB): A preview showing the first 100 lines with file metadata
            - For documents: Extracted text content
        """
        logger.debug(
            "read_file called with file_path: %s, workspace_id: %s",
            file_path,
            self.workspace.id,
        )

        # Smart search: first look in input directory, then in output directory
        resolved_path = self.workspace.resolve_path_with_search(file_path)
        logger.debug("Resolved path: %s", resolved_path)

        # Simple retry mechanism for potential timing issues
        max_retries = 3
        retry_delay = 0.1  # 100ms

        for attempt in range(max_retries):
            if resolved_path.exists():
                logger.debug("File found on attempt %d", attempt + 1)
                break

            logger.debug(
                "File not found on attempt %d, retrying in %.3fs",
                attempt + 1,
                retry_delay,
            )

            if attempt == max_retries - 1:
                # Last attempt failed, raise detailed error
                workspace_dirs = [
                    str(self.workspace.input_dir),
                    str(self.workspace.output_dir),
                    str(self.workspace.temp_dir),
                ]
                error_msg = (
                    f"File not found: {file_path} "
                    f"(search path: {resolved_path}, "
                    f"workspace directories: {workspace_dirs}, "
                    f"retry attempts: {max_retries})"
                )
                logger.error("FileNotFoundError: %s", error_msg)
                raise FileNotFoundError(error_msg)

            time.sleep(retry_delay)

        logger.debug("Reading file: %s", resolved_path)

        # Step 1: Check if this is a document file that requires special parsing
        # (Document files bypass binary checks and go directly to parser)
        if is_document_file(str(resolved_path)):
            logger.debug("Detected document file, using document parser")
            content = extract_text_from_document(str(resolved_path))
            logger.debug(
                "Successfully extracted %d characters from document %s",
                len(content),
                resolved_path,
            )
            return content

        # Step 2: Quick MIME check (no file I/O)
        if _is_binary_by_mime(str(resolved_path)):
            raise ValueError(
                f"Cannot read binary file: {file_path}. "
                f"This tool only supports text files."
            )

        # Step 3: Open file ONCE in binary mode for all checks and reading
        with open(resolved_path, "rb") as f:
            # Read first chunk for safety checks
            chunk = f.read(8192)

            # Binary file detection (MIME + content checks)
            if _is_binary_file(str(resolved_path), chunk):
                raise ValueError(
                    f"Cannot read binary file: {file_path}. "
                    f"This tool only supports text files."
                )

            # Get file size
            file_size = os.path.getsize(resolved_path)

            # Check if file exceeds size limit
            if file_size > DEFAULT_MAX_FILE_SIZE:
                # File is too large - generate preview without reading entire file
                logger.debug("File too large (%d bytes), generating preview", file_size)
                return _get_file_preview_from_file(
                    f, file_size, encoding, PREVIEW_LINES
                )

            # File is within size limit - read and decode full content
            f.seek(0)  # Seek back to beginning
            content = f.read().decode(encoding)
            logger.debug(
                "Successfully read %d bytes from %s", len(content), resolved_path
            )
            return content

    def write_file(
        self,
        file_path: str,
        content: str,
        encoding: str = "utf-8",
        create_dirs: bool = True,
    ) -> bool:
        """Write file content in workspace"""
        logger.debug(
            "write_file called with file_path: %s, content_length: %d, workspace_id: %s",
            file_path,
            len(content),
            self.workspace.id,
        )

        resolved_path = self._resolve_path(file_path, "output")
        logger.debug("Resolved path: %s", resolved_path)

        if create_dirs:
            logger.debug("Creating parent directories for: %s", resolved_path.parent)
            resolved_path.parent.mkdir(parents=True, exist_ok=True)

        logger.debug("Writing %d bytes to: %s", len(content), resolved_path)
        with open(resolved_path, "w", encoding=encoding) as f:
            f.write(content)

        logger.debug("Successfully wrote file: %s", resolved_path)
        return True

    def append_file(
        self,
        file_path: str,
        content: str,
        encoding: str = "utf-8",
        create_dirs: bool = True,
    ) -> bool:
        """Append content to file in workspace"""
        resolved_path = self.workspace.resolve_path_with_search(file_path)

        if create_dirs:
            resolved_path.parent.mkdir(parents=True, exist_ok=True)

        with open(resolved_path, "a", encoding=encoding) as f:
            f.write(content)
        return True

    def delete_file(self, file_path: str) -> bool:
        """Delete file in workspace"""
        resolved_path = self.workspace.resolve_path_with_search(file_path)

        if not resolved_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        resolved_path.unlink()
        return True

    def file_exists(self, file_path: str) -> bool:
        """Check if file exists in workspace"""
        try:
            resolved_path = self.workspace.resolve_path_with_search(file_path)
            return resolved_path.exists()
        except (ValueError, FileNotFoundError):
            # ValueError: path is outside allowed directories
            # FileNotFoundError: file not found in searched directories
            return False

    def list_files(
        self,
        directory_path: str = ".",
        show_hidden: bool = False,
        recursive: bool = False,
    ) -> Dict[str, Any]:
        """List files in workspace directory (default: list all directories)"""
        # If no directory path specified, return all directories' files
        if directory_path == ".":
            return self.workspace.get_all_files()

        # If specific directory is specified, only list files in that directory
        resolved_path = self._resolve_path(directory_path)

        if not resolved_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")

        files = []

        def scan_directory(current_path: Path, is_root: bool = True) -> None:
            try:
                for item in current_path.iterdir():
                    if not show_hidden and item.name.startswith("."):
                        continue

                    stat = item.stat()
                    file_info = FileInfo(
                        name=item.name,
                        path=str(item),
                        size=stat.st_size,
                        is_file=item.is_file(),
                        is_dir=item.is_dir(),
                        modified_time=stat.st_mtime,
                    )
                    files.append(file_info)

                    if recursive and item.is_dir():
                        scan_directory(item, False)

            except PermissionError:
                pass

        scan_directory(resolved_path)

        return {
            "files": [file.dict() for file in files],
            "total_count": len(files),
            "current_path": str(resolved_path),
            "directory": directory_path,
        }

    def create_directory(self, directory_path: str, parents: bool = True) -> bool:
        """Create directory in workspace"""
        resolved_path = self._resolve_path(directory_path)
        resolved_path.mkdir(parents=parents, exist_ok=True)
        return True

    def get_file_info(self, file_path: str) -> FileInfo:
        """Get detailed information about file in workspace"""
        resolved_path = self.workspace.resolve_path_with_search(file_path)

        if not resolved_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        stat = resolved_path.stat()

        return FileInfo(
            name=resolved_path.name,
            path=str(resolved_path),
            size=stat.st_size,
            is_file=resolved_path.is_file(),
            is_dir=resolved_path.is_dir(),
            modified_time=stat.st_mtime,
        )

    def read_json_file(self, file_path: str, encoding: str = "utf-8") -> Any:
        """Read JSON file in workspace"""
        from .file_tool import read_json_file as basic_read_json_file

        resolved_path = self.workspace.resolve_path_with_search(file_path)
        return basic_read_json_file(str(resolved_path), encoding)

    def write_json_file(
        self,
        file_path: str,
        data: Dict[str, Any],
        encoding: str = "utf-8",
        indent: int = 2,
    ) -> bool:
        """Write JSON file in workspace"""
        from .file_tool import write_json_file as basic_write_json_file

        resolved_path = self._resolve_path(file_path, "output")
        return basic_write_json_file(str(resolved_path), data, encoding, indent)

    def read_csv_file(
        self,
        file_path: str,
        encoding: str = "utf-8",
        delimiter: str = ",",
    ) -> List[Dict[str, str]]:
        """Read CSV file in workspace"""
        # Read the file directly without using document parser
        # CSV files should be read as plain text for proper parsing
        resolved_path = self.workspace.resolve_path_with_search(file_path)

        if not resolved_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        with open(resolved_path, "r", encoding=encoding) as f:
            content = f.read()

        lines = content.strip().split("\n")
        if not lines:
            return []

        reader = csv.DictReader(lines, delimiter=delimiter)
        return list(reader)

    def write_csv_file(
        self,
        file_path: str,
        data: List[Dict[str, str]],
        encoding: str = "utf-8",
        delimiter: str = ",",
    ) -> bool:
        """Write CSV file in workspace"""
        from .file_tool import write_csv_file as basic_write_csv_file

        resolved_path = self._resolve_path(file_path, "output")
        return basic_write_csv_file(str(resolved_path), data, encoding, delimiter)

    def get_workspace_output_files(self) -> Dict[str, Any]:
        """Get output file list from current workspace"""
        try:
            output_files = self.workspace.get_output_files()

            return {
                "success": True,
                "id": self.workspace.id,
                "workspace_dir": str(self.workspace.workspace_dir),
                "output_dir": str(self.workspace.output_dir),
                "file_count": len(output_files),
                "files": output_files,
            }
        except Exception as e:
            return {"error": str(e), "files": []}

    def edit_file(
        self,
        file_path: str,
        operations: List[Union[Dict[str, Any], EditOperation]],
        encoding: str = "utf-8",
        backup: bool = False,
    ) -> EditResult:
        """Precise file editing in workspace"""
        logger.debug(
            "edit_file called with file_path: %s, operations_count: %d, workspace_id: %s",
            file_path,
            len(operations),
            self.workspace.id,
        )

        # Import the edit_file function from the basic file_tool module
        from .file_tool import edit_file as basic_edit_file

        # Resolve the file path within the workspace
        resolved_path = self._resolve_path_with_search(file_path)
        logger.debug("Resolved path: %s", resolved_path)

        # Convert to string path for the basic edit_file function
        str_path = str(resolved_path)

        # Call the basic edit_file function with the resolved path
        result = basic_edit_file(str_path, operations, encoding, backup)

        logger.debug("edit_file result: %s", result)
        return result

    def find_and_replace(
        self,
        file_path: str,
        pattern: str,
        replacement: str,
        encoding: str = "utf-8",
        use_regex: bool = False,
        case_sensitive: bool = False,
        backup: bool = False,
    ) -> EditResult:
        """Find and replace text content in workspace"""
        logger.debug(
            "find_and_replace called with file_path: %s, pattern: %s, workspace_id: %s",
            file_path,
            pattern,
            self.workspace.id,
        )

        # Import the find_and_replace function from the basic file_tool module
        from .file_tool import find_and_replace as basic_find_and_replace

        # Resolve the file path within the workspace
        resolved_path = self._resolve_path_with_search(file_path)
        logger.debug("Resolved path: %s", resolved_path)

        # Convert to string path for the basic find_and_replace function
        str_path = str(resolved_path)

        # Call the basic find_and_replace function with the resolved path
        result = basic_find_and_replace(
            str_path, pattern, replacement, encoding, use_regex, case_sensitive, backup
        )

        logger.debug("find_and_replace result: %s", result)
        return result

    def _resolve_path_with_search(self, file_path: str) -> Path:
        """Intelligently resolve file path in workspace (first in input directory, then in output directory)"""
        logger.debug("_resolve_path_with_search called with file_path: %s", file_path)

        # Use the centralized workspace method
        return self.workspace.resolve_path_with_search(file_path)

    def _resolve_path(self, file_path: str, default_dir: str = "output") -> Path:
        """Resolve file path within workspace"""
        logger.debug(
            "_resolve_path called with file_path: %s, default_dir: %s",
            file_path,
            default_dir,
        )

        path = Path(file_path)

        if path.is_absolute():
            # Absolute path: check if within workspace
            abs_path = path.resolve()
            workspace_abs = self.workspace.workspace_dir.resolve()

            logger.debug(
                "Absolute path check - abs_path: %s, workspace_abs: %s",
                abs_path,
                workspace_abs,
            )

            if abs_path == workspace_abs or abs_path.is_relative_to(workspace_abs):
                logger.debug("Absolute path resolved to: %s", abs_path)
                return abs_path

            # Special case: if the path starts with "/workspace", provide a helpful error message
            if str(abs_path).startswith("/workspace"):
                error_msg = (
                    f"Path '{file_path}' is not within the workspace. "
                    f"Please use relative paths instead of absolute paths. For example, use 'filename.txt' instead of '/workspace/filename.txt'. "
                    f"Current workspace directory: {workspace_abs}"
                )
            else:
                error_msg = f"Path '{file_path}' is not within the workspace"

            logger.error("ValueError: %s", error_msg)
            raise ValueError(error_msg)
        else:
            # Relative path: resolve to specified directory
            # Special handling: if path already starts with output/input/temp, use workspace root
            path_str = str(path)
            first_component = (
                path_str.split("/")[0]
                if "/" in path_str
                else path_str.split("\\")[0]
                if "\\" in path_str
                else ""
            )

            if first_component in ["output", "input", "temp"]:
                # Path already includes the directory prefix, resolve to workspace root
                resolved_path = (self.workspace.workspace_dir / path).resolve()
            elif default_dir == "input":
                resolved_path = (self.workspace.input_dir / path).resolve()
            elif default_dir == "output":
                resolved_path = (self.workspace.output_dir / path).resolve()
            elif default_dir == "temp":
                resolved_path = (self.workspace.temp_dir / path).resolve()
            else:
                resolved_path = (self.workspace.workspace_dir / path).resolve()

            logger.debug("Relative path resolved to: %s", resolved_path)
            return resolved_path


def _get_workspace_ops(workspace_id: str) -> WorkspaceFileOperations:
    """Helper function to initialize WorkspaceFileOperations."""
    workspace = TaskWorkspace(id=workspace_id)
    return WorkspaceFileOperations(workspace)


def workspace_read_file(
    workspace_id: str, file_path: str, encoding: str = "utf-8"
) -> str:
    """
    Reads the content of a file within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file relative to the workspace.
        encoding: The encoding to use for reading the file (default: 'utf-8').

    Returns:
        The content of the file as a string.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.read_file(file_path, encoding)


def workspace_write_file(
    workspace_id: str,
    file_path: str,
    content: str,
    encoding: str = "utf-8",
    create_dirs: bool = True,
) -> bool:
    """
    Writes content to a file within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file relative to the workspace.
        content: The content to write to the file.
        encoding: The encoding to use for writing the file (default: 'utf-8').
        create_dirs: If True, creates parent directories if they don't exist (default: True).

    Returns:
        True if the write operation was successful.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.write_file(file_path, content, encoding, create_dirs)


def workspace_list_files(
    workspace_id: str,
    directory: str = ".",
    recursive: bool = False,
    show_hidden: bool = False,
) -> dict[str, Any]:
    """
    Lists files and directories within a specified path in the workspace.

    Args:
        workspace_id: The ID of the workspace.
        directory: The directory path to list (default: '.').
        recursive: If True, lists files recursively (default: False).
        show_hidden: If True, includes hidden files (default: False).

    Returns:
        A dictionary containing file information, total count, and path details.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.list_files(directory, show_hidden, recursive)


def workspace_file_exists(workspace_id: str, file_path: str) -> bool:
    """
    Checks if a file or directory exists within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file or directory relative to the workspace.

    Returns:
        True if the file or directory exists, False otherwise.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.file_exists(file_path)


def workspace_delete_file(workspace_id: str, file_path: str) -> bool:
    """
    Deletes a file within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file relative to the workspace.

    Returns:
        True if the file was successfully deleted.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.delete_file(file_path)


def workspace_create_directory(
    workspace_id: str, directory_path: str, parents: bool = True
) -> bool:
    """
    Creates a directory within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        directory_path: The path of the directory to create.
        parents: If True, creates any necessary intermediate directories (default: True).

    Returns:
        True if the directory was successfully created.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.create_directory(directory_path, parents)


def workspace_get_file_info(workspace_id: str, file_path: str) -> FileInfo:
    """
    Retrieves detailed information about a file or directory in the workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file or directory relative to the workspace.

    Returns:
        A FileInfo object containing metadata about the file.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.get_file_info(file_path)


def workspace_read_json_file(
    workspace_id: str, file_path: str, encoding: str = "utf-8"
) -> Any:
    """
    Reads and parses a JSON file within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the JSON file relative to the workspace.
        encoding: The encoding to use for reading the file (default: 'utf-8').

    Returns:
        The parsed JSON content (Any type).
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.read_json_file(file_path, encoding)


def workspace_write_json_file(
    workspace_id: str,
    file_path: str,
    data: Dict[str, Any],
    encoding: str = "utf-8",
    indent: int = 2,
) -> bool:
    """
    Writes data as JSON to a file within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the JSON file relative to the workspace.
        data: The dictionary data to write.
        encoding: The encoding to use for writing the file (default: 'utf-8').
        indent: The indentation level for the JSON output (default: 2).

    Returns:
        True if the write operation was successful.
    """
    ops = _get_workspace_ops(workspace_id)
    # Note: The original ops.write_json_file signature uses 'indent', not 'create_dirs' as the 4th arg.
    # We rely on ops.write_file's default behavior for create_dirs or adjust the call if necessary.
    # Assuming the internal implementation handles directory creation via write_file.
    return ops.write_json_file(file_path, data, encoding, indent)


def workspace_read_csv_file(
    workspace_id: str, file_path: str, encoding: str = "utf-8", delimiter: str = ","
) -> List[Dict[str, str]]:
    """
    Reads and parses a CSV file within the specified workspace, returning a list of dictionaries.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the CSV file relative to the workspace.
        encoding: The encoding to use for reading the file (default: 'utf-8').
        delimiter: The delimiter character used in the CSV file (default: ',').

    Returns:
        A list of dictionaries, where each dictionary represents a row.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.read_csv_file(file_path, encoding, delimiter)


def workspace_write_csv_file(
    workspace_id: str,
    file_path: str,
    data: List[Dict[str, str]],
    encoding: str = "utf-8",
    delimiter: str = ",",
) -> bool:
    """
    Writes a list of dictionaries as CSV content to a file within the specified workspace.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the CSV file relative to the workspace.
        data: A list of dictionaries representing the CSV rows.
        encoding: The encoding to use for writing the file (default: 'utf-8').
        delimiter: The delimiter character to use in the CSV file (default: ',').

    Returns:
        True if the write operation was successful.
    """
    ops = _get_workspace_ops(workspace_id)
    # The original ops.write_csv_file signature uses List[Dict[str, str]], not List[List[Any]].
    # Also, the original signature takes 'delimiter' instead of 'create_dirs' as the last positional arg.
    return ops.write_csv_file(file_path, data, encoding, delimiter)


def workspace_edit_file(
    workspace_id: str,
    file_path: str,
    operations: List[Union[Dict[str, Any], EditOperation]],
    encoding: str = "utf-8",
    backup: bool = False,
) -> EditResult:
    """
    Performs precise line-based edits (insert, delete, replace) on a file.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file relative to the workspace.
        operations: A list of edit operations (dictionaries or EditOperation objects).
        encoding: The encoding of the file (default: 'utf-8').
        backup: If True, creates a backup of the original file (default: False).

    Returns:
        An EditResult object detailing the result of the operation.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.edit_file(file_path, operations, encoding, backup)


def workspace_find_and_replace(
    workspace_id: str,
    file_path: str,
    pattern: str,
    replacement: str,
    encoding: str = "utf-8",
    use_regex: bool = False,
    case_sensitive: bool = False,
    backup: bool = False,
) -> EditResult:
    """
    Finds occurrences of a pattern in a file and replaces them with the replacement string.

    Args:
        workspace_id: The ID of the workspace.
        file_path: The path to the file relative to the workspace.
        pattern: The string or regex pattern to search for.
        replacement: The string to replace the pattern with.
        encoding: The encoding of the file (default: 'utf-8').
        use_regex: If True, treats the pattern as a regular expression (default: False).
        case_sensitive: If True, performs a case-sensitive search (default: False).
        backup: If True, creates a backup of the original file (default: False).

    Returns:
        An EditResult object detailing the result of the operation.
    """
    ops = _get_workspace_ops(workspace_id)
    return ops.find_and_replace(
        file_path, pattern, replacement, encoding, use_regex, case_sensitive, backup
    )
